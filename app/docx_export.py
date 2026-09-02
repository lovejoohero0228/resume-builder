"""이력서를 Resume 1 서식(Arial · 남색 #153D63 · 섹션 대문자+하단선 · 회사 볼드)의
진짜 .docx 로 생성. Word 네이티브 서식이라 Word→PDF 가 미리보기와 일관됨.

resume_struct(dict) -> docx bytes
"""
from __future__ import annotations
import io
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.enum.text import WD_TAB_ALIGNMENT, WD_ALIGN_PARAGRAPH

# ---- CoverLetter2 서식 스펙 (Calibri · 남색 #1F3864 · 중앙 헤더 · 좁은 여백) ----
FONT = "Calibri"
EA = "Malgun Gothic"      # 한글 fallback
NAVY = RGBColor(0x1F, 0x38, 0x64)
GRAY_DATE = RGBColor(0x59, 0x59, 0x59)   # 기간·연도
GRAY_DESC = RGBColor(0x40, 0x40, 0x40)   # 회사 아래 회색 설명
BLACK = RGBColor(0x00, 0x00, 0x00)
BODY = 9.5
CENTER = WD_ALIGN_PARAGRAPH.CENTER


def _run(p, text, color=BLACK, bold=False, size=BODY):
    r = p.add_run(text)
    f = r.font
    f.name = FONT
    f.size = Pt(size)
    f.color.rgb = color
    r.bold = bold
    rpr = r._element.get_or_add_rPr()
    rf = rpr.get_or_add_rFonts()
    rf.set(qn("w:ascii"), FONT)
    rf.set(qn("w:hAnsi"), FONT)
    rf.set(qn("w:cs"), FONT)
    rf.set(qn("w:eastAsia"), EA)
    return r


def _rich(p, text, color=BLACK, bold=False, size=BODY):
    """`**bold**` 토글을 반영해 여러 run 으로 추가."""
    for i, seg in enumerate(str(text).split("**")):
        if seg:
            _run(p, seg, color, bold or (i % 2 == 1), size)


def _para(doc, before=0.0, after=1.0, align=None):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = 1.04
    if align is not None:
        p.alignment = align
    return p


def _section(doc, title):
    p = _para(doc, before=4.6, after=1.9)
    t = str(title)
    _run(p, t.upper() if t.isascii() else t, NAVY, True, 10)


def resume_docx_bytes(rs: dict) -> bytes:
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(BODY)
    rpr = normal.element.get_or_add_rPr()
    rf = rpr.get_or_add_rFonts()
    rf.set(qn("w:eastAsia"), EA)

    sec = doc.sections[0]
    sec.top_margin = Inches(0.28)
    sec.bottom_margin = Inches(0.14)
    sec.left_margin = Inches(0.39)
    sec.right_margin = Inches(0.39)
    right_tab = Inches(8.5 - 0.39 - 0.39)   # 콘텐츠 우측 끝 (기간·연도 우측 정렬)

    def _right_tab(p):
        p.paragraph_format.tab_stops.add_tab_stop(right_tab, WD_TAB_ALIGNMENT.RIGHT)

    lang = rs.get("lang", "en")
    ko = lang == "ko"
    labels = ({"p": "프로필", "s": "기술 스택", "e": "경력"} if ko
              else {"p": "PROFILE", "s": "SKILLS", "e": "EXPERIENCE"})

    def block_identity():
        idn = rs.get("identity") or {}
        if not idn:
            return
        _run(_para(doc, after=1.5, align=CENTER), idn.get("name", ""), NAVY, True, 16)
        if idn.get("tagline"):
            _rich(_para(doc, after=1.5, align=CENTER), idn["tagline"], BLACK, True, 10)
        if idn.get("contact"):
            _run(_para(doc, after=2, align=CENTER), idn["contact"], BLACK, False, 9)
        if idn.get("summary"):
            _section(doc, labels["p"])
            _rich(_para(doc, after=1), idn["summary"], BLACK, False, BODY)

    def block_skills():
        vis = [s for s in (rs.get("skills") or []) if not s.get("hidden")]
        if not vis:
            return
        _section(doc, labels["s"])
        for s in vis:
            p = _para(doc, before=0.9, after=0.9)
            if s.get("label"):
                _run(p, s["label"] + "   ", BLACK, True, BODY)
            _rich(p, s.get("rest", ""), BLACK, False, BODY)

    def block_education():
        for g in rs.get("education", []) or []:
            _section(doc, g.get("title", ""))
            for it in g.get("items", []):
                p = _para(doc, before=0.4, after=0.6)
                if it.get("sub"):
                    p.paragraph_format.left_indent = Inches(0.18)
                _rich(p, it.get("text", ""), GRAY_DESC if it.get("sub") else BLACK, False, BODY)

    exps = list(rs.get("experiences") or [])
    _exp_state = {"i": 0, "header": False}

    def block_experience():
        if _exp_state["i"] >= len(exps):
            return
        if not _exp_state["header"]:
            _section(doc, labels["e"])
            _exp_state["header"] = True
        e = exps[_exp_state["i"]]
        _exp_state["i"] += 1
        # 1줄: **회사 | 직함**  [우측탭] 기간 · 위치
        p = _para(doc, before=4.25, after=0)
        _right_tab(p)
        head = e.get("company", "")
        if e.get("title"):
            head += "  |  " + e["title"]
        _run(p, head, BLACK, True, 10)
        when = "  ·  ".join(x for x in [e.get("period"), e.get("location")] if x)
        if when:
            _run(p, "\t", BLACK, False, 9)
            _run(p, when, GRAY_DATE, False, 9)
        # 2줄: 회색 설명(컨텍스트)
        if e.get("context"):
            _rich(_para(doc, after=2), e["context"], GRAY_DESC, False, 9)
        # 불릿 (볼드 리드)
        for b in e.get("bullets", []):
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_before = Pt(0.55)
            p.paragraph_format.space_after = Pt(0.55)
            p.paragraph_format.line_spacing = 1.04
            _rich(p, b, BLACK, False, BODY)

    blocks = {"identity": block_identity, "skills": block_skills,
              "experience": block_experience, "education": block_education}
    order = rs.get("order") or ["identity", "skills", "experience", "education"]
    if "experience" not in order and exps:
        order = order + ["experience"] * len(exps)
    for tok in order:
        fn = blocks.get(tok)
        if fn:
            fn()

    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()
